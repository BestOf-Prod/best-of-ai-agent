import re
import logging
from docx import Document
from typing import List
from utils.logger import setup_logging

# Setup logging
logger = setup_logging(__name__)

def extract_urls_from_docx(file_content) -> List[str]:
    """
    Extract URLs from a Word document (.docx file)
    
    Args:
        file_content: The uploaded file content from Streamlit
        
    Returns:
        List[str]: List of unique URLs found in the document
    """
    logger.info("Starting URL extraction from Word document")
    
    try:
        # Load the document from the file content
        doc = Document(file_content)
        logger.info(f"Successfully loaded Word document")
        
        urls = []  # Use list to preserve order
        seen_urls = set()  # Track duplicates
        
        # Extract URLs from paragraph text in document order
        logger.debug("Extracting URLs from paragraph text")
        for paragraph in doc.paragraphs:
            text = paragraph.text
            found_urls = extract_urls_from_text(text)
            # Add URLs in the order they appear, avoiding duplicates
            for url in found_urls:
                if url not in seen_urls:
                    urls.append(url)
                    seen_urls.add(url)
        
        # Extract URLs from hyperlinks in document order
        logger.debug("Extracting URLs from hyperlinks")
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.element.get("{http://www.w3.org/XML/1998/namespace}space") is None:
                    # Check if the run contains a hyperlink
                    for hyperlink in run.element.xpath('.//w:hyperlink'):
                        r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if r_id:
                            try:
                                url = doc.part.rels[r_id].target_ref
                                if url and is_valid_url(url) and url not in seen_urls:
                                    urls.append(url)
                                    seen_urls.add(url)
                                    logger.debug(f"Found hyperlink URL: {url}")
                            except Exception as e:
                                logger.warning(f"Error processing hyperlink: {str(e)}")
        
        # Extract URLs from tables in document order
        logger.debug("Extracting URLs from tables")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found_urls = extract_urls_from_text(cell.text)
                    # Add URLs in the order they appear in table cells
                    for url in found_urls:
                        if url not in seen_urls:
                            urls.append(url)
                            seen_urls.add(url)
        
        url_list = urls  # Already a list in correct order
        logger.info(f"Successfully extracted {len(url_list)} unique URLs from document in document order")
        
        # Log first few URLs to verify order preservation
        if url_list:
            sample_urls = url_list[:3]  # First 3 URLs
            logger.info(f"URLs extracted in order (showing first 3): {sample_urls}")
        
        return url_list
        
    except Exception as e:
        logger.error(f"Error extracting URLs from Word document: {str(e)}")
        raise Exception(f"Failed to process Word document: {str(e)}")

def extract_urls_from_text(text: str) -> List[str]:
    """
    Extract URLs from plain text using regex, preserving order
    
    Args:
        text (str): The text to search for URLs
        
    Returns:
        List[str]: List of URLs found in the text in order of appearance
    """
    if not text:
        return []
    
    # Regex pattern to match URLs
    url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+(?:[^\s<>"{}|\\^`\[\].,;!?])'
    
    urls = []
    seen_urls = set()
    matches = re.finditer(url_pattern, text, re.IGNORECASE)
    
    for match in matches:
        url = match.group(0)
        if is_valid_url(url) and url not in seen_urls:
            urls.append(url)
            seen_urls.add(url)
    
    return urls

def is_valid_url(url: str) -> bool:
    """
    Validate if a URL is properly formatted and not a local/relative URL
    
    Args:
        url (str): The URL to validate
        
    Returns:
        bool: True if the URL is valid, False otherwise
    """
    if not url:
        return False
    
    # Basic validation - must start with http or https
    if not url.startswith(('http://', 'https://')):
        return False
    
    # Must contain a domain
    if '.' not in url:
        return False
    
    # Filter out obviously invalid URLs
    invalid_patterns = [
        'localhost',
        '127.0.0.1',
        '192.168.',
        '10.0.',
        'file://',
        'ftp://',
    ]
    
    for pattern in invalid_patterns:
        if pattern in url.lower():
            return False
    
    return True

def validate_document_format(filename: str) -> bool:
    """
    Validate that the uploaded file is a Word document
    
    Args:
        filename (str): The name of the uploaded file
        
    Returns:
        bool: True if it's a valid Word document format
    """
    if not filename:
        return False
    
    valid_extensions = ['.docx', '.doc']
    return any(filename.lower().endswith(ext) for ext in valid_extensions) 