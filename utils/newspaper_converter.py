import os
import re
import tempfile
import requests
import zipfile
import io
import random
import time
from urllib.parse import urlparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from bs4 import BeautifulSoup
import logging
from datetime import datetime
from utils.logger import setup_logging
from utils.google_drive_manager import GoogleDriveManager

# Import capsule parser for typography specifications
from utils.capsule_parser import get_typography_for_article, TypographySpec

logger = setup_logging(__name__, log_level=logging.INFO)

# Font matrix for different site types
FONT_MATRIX = {
    'newspaper_sites': {
        'serif_fonts': [
            'Times New Roman',
            'Georgia',
            'Playfair Display',
            'Minion Pro',
            'Adobe Garamond Pro',
            'Times',
            'Book Antiqua',
            'Century Schoolbook'
        ],
        'domains': [
            'nytimes.com',
            'washingtonpost.com',
            'wsj.com',
            'theguardian.com',
            'ft.com',
            'latimes.com',
            'chicagotribune.com',
            'usatoday.com',
            'boston.com',
            'sfgate.com'
        ]
    },
    'web_news_sites': {
        'sans_serif_fonts': [
            'Arial',
            'Helvetica',
            'Calibri',
            'Verdana',
            'Tahoma',
            'Franklin Gothic Medium',
            'Segoe UI',
            'Roboto'
        ],
        'domains': [
            'espn.com',
            'cnn.com',
            'bbc.com',
            'reuters.com',
            'buzzfeed.com',
            'techcrunch.com',
            'mashable.com',
            'theverge.com',
            'wired.com',
            'engadget.com'
        ]
    }
}

def get_font_for_site(url):
    """
    Determine the appropriate font based on the site URL
    
    Args:
        url (str): The source URL of the article
        
    Returns:
        str: Font name to use for the article
    """
    if not url:
        return random.choice(FONT_MATRIX['newspaper_sites']['serif_fonts'])
    
    domain = urlparse(url).netloc.replace('www.', '').lower()
    
    # Check if it's a newspaper site
    if any(news_domain in domain for news_domain in FONT_MATRIX['newspaper_sites']['domains']):
        return random.choice(FONT_MATRIX['newspaper_sites']['serif_fonts'])
    
    # Check if it's a web news site
    if any(web_domain in domain for web_domain in FONT_MATRIX['web_news_sites']['domains']):
        return random.choice(FONT_MATRIX['web_news_sites']['sans_serif_fonts'])
    
    # Default fallback - assume newspaper for unknown sites
    return random.choice(FONT_MATRIX['newspaper_sites']['serif_fonts'])

def calculate_word_count(text):
    """
    Calculate word count for article text
    
    Args:
        text (str): The text to count words in
        
    Returns:
        int: Number of words
    """
    if not text:
        return 0
    
    # Split by whitespace and filter out empty strings
    words = [word.strip() for word in text.split() if word.strip()]
    return len(words)

def create_directory_name(article_data, font_name, word_count):
    """
    Create directory name with word count, font, and article title
    
    Args:
        article_data (dict): Article data dictionary
        font_name (str): Font used for the article
        word_count (int): Word count of the article body
        
    Returns:
        str: Directory name
    """
    title = article_data.get('headline', 'Unknown Title')
    # Clean title for directory name
    clean_title = sanitize_filename(title)
    
    # Create directory name: WordCount_Font_Title
    font_clean = font_name.replace(' ', '-')
    directory_name = f"{word_count}words_{font_clean}_{clean_title}"
    
    # Ensure directory name isn't too long (limit to 100 chars)
    if len(directory_name) > 100:
        max_title_length = 100 - len(f"{word_count}words_{font_clean}_")
        clean_title = clean_title[:max_title_length]
        directory_name = f"{word_count}words_{font_clean}_{clean_title}"
    
    return directory_name

def sanitize_filename(title):
    """Convert article title to safe filename with hyphens"""
    if not title:
        return "untitled-article"
    
    # Remove or replace problematic characters
    title = title.strip()
    
    # Replace spaces and common punctuation with hyphens
    title = re.sub(r'[^\w\s-]', '', title)  # Remove special chars except word chars, spaces, hyphens
    title = re.sub(r'\s+', '-', title)      # Replace spaces with hyphens
    title = re.sub(r'-+', '-', title)       # Replace multiple hyphens with single hyphen
    title = title.strip('-')                # Remove leading/trailing hyphens
    
    # Convert to lowercase for consistency
    title = title.lower()
    
    # Limit length to avoid filesystem issues
    if len(title) > 50:
        # Find the last word boundary before 50 characters
        truncated = title[:50]
        if '-' in truncated:
            last_hyphen = truncated.rfind('-')
            title = truncated[:last_hyphen]
        else:
            title = truncated
    
    # Ensure we have something
    if not title:
        title = "untitled-article"
    
    return title

def extract_title_from_markdown(md_content):
    """Extract the first headline from markdown content"""
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        # Look for H1 or H2 headers
        if line.startswith('# '):
            return line[2:].strip()
        elif line.startswith('## '):
            return line[3:].strip()
    
    # Fallback: use first non-empty line
    for line in lines:
        line = line.strip()
        if line and not line.startswith('#'):
            # Take first sentence or up to 50 chars
            if '.' in line:
                return line.split('.')[0].strip()
            else:
                return line[:50].strip()
    
    return "Untitled Article"

def extract_images_from_markdown(md_content):
    """Extract image URLs and alt text from markdown content."""
    image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    images = []
    
    logger.info(f"Searching for images in markdown content (length: {len(md_content)})")
    
    for match in re.finditer(image_pattern, md_content):
        alt_text = match.group(1)
        url = match.group(2)
        images.append({
            'alt_text': alt_text,
            'url': url,
            'markdown': match.group(0)
        })
        logger.info(f"Found image: {url} (alt: {alt_text})")
    
    logger.info(f"Total images found: {len(images)}")
    return images

def download_image(url, temp_dir):
    """Download image from URL and save to temporary directory."""
    logger.info(f"Attempting to download image from: {url}")
    
    try:
        # Validate URL
        if not url or not url.startswith(('http://', 'https://')):
            logger.error(f"Invalid URL format: {url}")
            return None
            
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        logger.info(f"Making request to {url}")
        response = requests.get(url, headers=headers, timeout=15)  # Increased timeout
        response.raise_for_status()
        
        # Check if response contains image data
        content_type = response.headers.get('content-type', '').lower()
        if not content_type.startswith('image/'):
            logger.warning(f"URL does not return image content. Content-Type: {content_type}")
            return None
        
        # Get file extension from URL or content type
        parsed_url = urlparse(url)
        filename = os.path.basename(parsed_url.path)
        if not filename or '.' not in filename:            
            if 'jpeg' in content_type or 'jpg' in content_type:
                filename = 'image.jpg'
            elif 'png' in content_type:
                filename = 'image.png'
            elif 'gif' in content_type:
                filename = 'image.gif'
            elif 'webp' in content_type:
                filename = 'image.webp'
            else:
                filename = 'image.jpg'
        
        temp_path = os.path.join(temp_dir, filename)
        
        # Write image data
        with open(temp_path, 'wb') as f:
            f.write(response.content)
        
        # Validate image dimensions after download
        try:
            from PIL import Image
            with Image.open(temp_path) as img:
                width, height = img.size
                
                # Skip very small images (likely icons or logos)
                if width < 200 or height < 200:
                    logger.info(f"Skipping small image ({width}x{height}): {temp_path}")
                    os.remove(temp_path)  # Clean up the file
                    return None
                
                # Skip very narrow or very wide images (likely banners or decorative elements)
                aspect_ratio = max(width, height) / min(width, height)
                if aspect_ratio > 4:  # More than 4:1 ratio
                    logger.info(f"Skipping unusual aspect ratio image ({width}x{height}, ratio {aspect_ratio:.1f}): {temp_path}")
                    os.remove(temp_path)  # Clean up the file
                    return None
                
                logger.info(f"Successfully downloaded and validated image: {temp_path} ({len(response.content)} bytes, {width}x{height})")
                
        except Exception as e:
            logger.warning(f"Could not validate image dimensions for {temp_path}: {str(e)}")
            # If we can't validate dimensions, keep the image but log the issue
        
        return temp_path
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Network error downloading image from {url}: {str(e)}")
        return None
    except Exception as e:
        logger.error(f"Unexpected error downloading image from {url}: {str(e)}")
        return None

def determine_layout(content_length):
    """Determine column layout based on content length."""
    if content_length < 1500:
        return 'single'  # Short: single column with headline above image, body below
    else:
        return 'double'  # Long: two columns with image on left

def add_column_break(doc):
    """Add a column break to the document."""
    p = doc.add_paragraph()
    run = p.runs[0] if p.runs else p.add_run()
    fldChar = OxmlElement('w:br')
    fldChar.set(qn('w:type'), 'column')
    run._element.append(fldChar)

def set_column_layout(section, num_columns):
    """Set the number of columns for a section."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_columns))
    if num_columns > 1:
        cols.set(qn('w:space'), '432')  # 0.3 inch spacing between columns

def apply_newspaper_styling(doc):
    """Apply newspaper-style formatting to the document."""
    # Set document margins (narrow margins for newspaper style)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

def create_headline_style(paragraph, text, font_name=None):
    """Apply headline styling to a paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = font_name or 'Times New Roman'
    run.font.size = Pt(38)  # Fixed headline size
    run.font.bold = True
    paragraph.space_after = Pt(6)

def create_capsule_based_style(paragraph, text, typography_spec: TypographySpec, alignment=WD_ALIGN_PARAGRAPH.LEFT, force_indent=True):
    """Apply styling based on capsule typography specifications."""
    logger.info(f"INDENT_DEBUG_STYLE: create_capsule_based_style called - force_indent={force_indent}, text='{text[:60]}...'")
    if not typography_spec:
        # Fallback to default styling
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return
    
    # Set alignment
    paragraph.alignment = alignment
    
    # Set indentation if specified in typography_spec OR if forced (for paragraph indentation)
    if typography_spec.indent > 0 or force_indent:
        indent_amount = typography_spec.indent if typography_spec.indent > 0 else 0.5
        paragraph.paragraph_format.first_line_indent = Inches(indent_amount)
        logger.info(f"INDENT_DEBUG_STYLE: Applied {indent_amount} inch first line indent (capsule-based)!")
    else:
        logger.info(f"INDENT_DEBUG_STYLE: No indentation applied (capsule-based)")
    
    # Create run and apply font specifications
    run = paragraph.add_run(text)
    run.font.name = typography_spec.font_family
    run.font.size = Pt(typography_spec.font_size)
    
    # Apply font weight
    if 'bold' in typography_spec.font_weight.lower():
        run.font.bold = True
    elif 'italic' in typography_spec.font_weight.lower():
        run.font.italic = True
    
    # Set line spacing (leading)
    paragraph.paragraph_format.line_spacing = Pt(typography_spec.leading)
    
    # Set spacing after paragraph
    paragraph.space_after = Pt(3)

def create_body_style(paragraph, text, font_name=None, indented=True):
    """Apply body text styling to a paragraph."""
    logger.info(f"INDENT_DEBUG_STYLE: create_body_style called - indented={indented}, text='{text[:60]}...'")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indented:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        logger.info(f"INDENT_DEBUG_STYLE: Applied 0.5 inch first line indent!")
    else:
        logger.info(f"INDENT_DEBUG_STYLE: No indentation applied")
    run = paragraph.add_run(text)
    run.font.name = font_name or 'Times New Roman'
    run.font.size = Pt(14)  # Fixed body text size
    paragraph.space_after = Pt(3)

def create_blockquote_style(paragraph, text, font_name=None):
    """Apply blockquote styling to a paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.right_indent = Inches(0.5)
    run = paragraph.add_run(text)
    run.font.name = font_name or 'Times New Roman'
    run.font.size = Pt(14)  # Fixed blockquote size to match body text
    run.font.italic = True
    paragraph.space_after = Pt(6)

def process_markdown_to_text(md_content):
    """Convert markdown to plain text and extract structure, preserving indentation."""
    logger.info(f"INDENT_DEBUG_CONVERTER_A: Input markdown content:\n{md_content[:500]}...")
    
    # First, let's try to preserve the structure by parsing the markdown directly
    # instead of converting to HTML first
    
    lines = md_content.split('\n')
    headline = None
    body_lines = []
    
    for line in lines:
        # Check for headlines
        if line.startswith('# ') and headline is None:
            headline = line[2:].strip()
        elif line.startswith('## ') and headline is None:
            headline = line[3:].strip()
        # Include all non-headline lines (including empty lines for paragraph breaks)
        elif not line.startswith('#'):
            body_lines.append(line)
    
    # Reconstruct body content preserving paragraph breaks and indentation
    body_content = '\n'.join(body_lines)
    
    # Split by double newlines to get paragraphs, but preserve indentation
    paragraphs = body_content.split('\n\n')
    cleaned_paragraphs = []
    
    for para in paragraphs:
        if para.strip():
            cleaned_paragraphs.append(para)
    
    body_text = '\n\n'.join(cleaned_paragraphs)
    
    logger.info(f"INDENT_DEBUG_CONVERTER_A: Output headline: '{headline}'")
    logger.info(f"INDENT_DEBUG_CONVERTER_A: Output body text:\n{body_text[:500]}...")
    
    return headline, body_text

def create_component_documents(article_data, temp_dir):
    """
    Create separate Word documents for each article component.
    
    Args:
        article_data: Dictionary containing article data with structured_content
        temp_dir: Temporary directory for file storage
        
    Returns:
        dict: Dictionary containing paths to component documents, images, and metadata
    """
    logger.info("Creating component documents from article data")
    
    # Get typography capsule or fallback to font selection
    typography_capsule = article_data.get('typography_capsule')
    font_name = None
    word_count = article_data.get('word_count', 0)
    
    if typography_capsule:
        logger.info(f"Using typography capsule {typography_capsule.capsule_id} for {word_count} words ({typography_capsule.category})")
        # Use the headline font as the primary font for directory naming
        headline_spec = typography_capsule.typography_specs.get('headline')
        if headline_spec:
            font_name = headline_spec.font_family
    else:
        # Fallback to original font selection method
        font_name = get_font_for_site(article_data.get('url', ''))
        logger.info(f"No capsule available, selected font: {font_name} for source: {article_data.get('source', 'unknown')}")
        
        # Calculate word count if not already provided
        if word_count == 0:
            article_text = article_data.get('text', '')
            word_count = calculate_word_count(article_text)
            logger.info(f"Calculated word count: {word_count}")
    
    # Create directory name with word count, font, and title
    directory_name = create_directory_name(article_data, font_name, word_count)
    logger.info(f"Directory name: {directory_name}")
    
    # Extract components
    headline = article_data.get('headline', '')
    structured_content = article_data.get('structured_content', [])
    
    # Get images from article data - handle both image_url and image_data
    images = []
    
    # Check for newspapers.com image_data (PIL Image object or bytes)
    if article_data.get('image_data'):
        # Save image data to temp directory for processing
        image_filename = f"newspapers_clipping_{int(time.time())}.png"
        image_path = os.path.join(temp_dir, image_filename)
        try:
            # Handle different image_data formats
            if hasattr(article_data['image_data'], 'save'):
                # It's a PIL Image object
                article_data['image_data'].save(image_path, 'PNG')
                logger.info(f"Saved PIL Image from newspapers.com: {image_path}")
            elif isinstance(article_data['image_data'], bytes):
                # It's raw bytes - convert to PIL Image first
                from PIL import Image
                import io
                img = Image.open(io.BytesIO(article_data['image_data']))
                img.save(image_path, 'PNG')
                logger.info(f"Converted bytes to PIL Image for newspapers.com: {image_path}")
            elif isinstance(article_data['image_data'], str):
                # It might be base64 encoded
                from PIL import Image
                import base64
                import io
                img_data = base64.b64decode(article_data['image_data'])
                img = Image.open(io.BytesIO(img_data))
                img.save(image_path, 'PNG')
                logger.info(f"Converted base64 to PIL Image for newspapers.com: {image_path}")
            else:
                logger.error(f"Unknown image_data format: {type(article_data['image_data'])}")
                raise ValueError(f"Unsupported image_data format: {type(article_data['image_data'])}")
            
            images.append({
                'alt_text': 'Newspapers.com Clipping',
                'url': 'newspapers.com_clipping',  # Special marker for clipping images
                'path': image_path,  # Direct path to image file
                'is_clipping': True
            })
            logger.info(f"Successfully processed image_data from newspapers.com: {image_path}")
        except Exception as e:
            logger.error(f"Failed to save newspapers.com image_data: {str(e)}")
            logger.error(f"Image data type: {type(article_data.get('image_data'))}")
            logger.error(f"Image data length: {len(article_data['image_data']) if hasattr(article_data['image_data'], '__len__') else 'N/A'}")
    
    # Check for regular image URL (from url_extractor)
    elif article_data.get('image_url'):
        images.append({
            'alt_text': 'Article Image',
            'url': article_data['image_url'],
            'markdown': f"![Article Image]({article_data['image_url']})",
            'is_clipping': False
        })
        logger.info(f"Using image URL from article data: {article_data['image_url']}")
    
    # Fallback: try to extract from markdown text
    else:
        extracted_images = extract_images_from_markdown(article_data.get('text', ''))
        for img in extracted_images:
            img['is_clipping'] = False
        images.extend(extracted_images)
        logger.info(f"Extracted {len(extracted_images)} images from markdown text")
    
    components = {
        'heading': None,
        'drophead': None,  # Subheading
        'author': None,
        'source': None, 
        'body': None,
        'blockquote': None,
        'caption': None,
        'photo_credit': None,
        'pullquote': None,
        'attribution': None,
        'image': None,
        'metadata': {
            'font_name': font_name,
            'word_count': word_count,
            'directory_name': directory_name,
            'title': article_data.get('headline', 'Unknown Title')
        }
    }
    
    # Create heading document
    if headline:
        heading_doc = Document()
        apply_newspaper_styling(heading_doc)
        set_column_layout(heading_doc.sections[0], 1)  # Single column
        
        heading_para = heading_doc.add_paragraph()
        
        # Use capsule-based styling if available
        if typography_capsule:
            headline_spec = typography_capsule.typography_specs.get('headline')
            if headline_spec:
                create_capsule_based_style(heading_para, headline, headline_spec, WD_ALIGN_PARAGRAPH.CENTER)
                logger.info(f"Applied capsule-based headline styling: {headline_spec.font_family} {headline_spec.font_size}pt")
            else:
                create_headline_style(heading_para, headline, font_name)
        else:
            create_headline_style(heading_para, headline, font_name)
        
        heading_path = os.path.join(temp_dir, 'heading.docx')
        heading_doc.save(heading_path)
        components['heading'] = heading_path
        logger.info(f"Created heading document: {heading_path}")
    
    # Create author document
    author = article_data.get('author')
    if author and author.strip() and author.strip() != "Unknown Author":
        author_doc = Document()
        apply_newspaper_styling(author_doc)
        set_column_layout(author_doc.sections[0], 1)
        
        author_para = author_doc.add_paragraph()
        if typography_capsule:
            author_spec = typography_capsule.typography_specs.get('author')
            if author_spec:
                create_capsule_based_style(author_para, f"By {author}", author_spec, WD_ALIGN_PARAGRAPH.LEFT)
                logger.info(f"Applied capsule-based author styling: {author_spec.font_family} {author_spec.font_size}pt")
            else:
                create_body_style(author_para, f"By {author}", font_name)
        else:
            create_body_style(author_para, f"By {author}", font_name)
        
        author_path = os.path.join(temp_dir, 'author.docx')
        author_doc.save(author_path)
        components['author'] = author_path
        logger.info(f"Created author document: {author_path}")
    
    # Create source document
    source = article_data.get('source')
    date = article_data.get('date')
    if source and source.strip() and source.strip() != "Unknown Source":
        source_doc = Document()
        apply_newspaper_styling(source_doc)
        set_column_layout(source_doc.sections[0], 1)
        
        source_para = source_doc.add_paragraph()
        source_text = f"{source}"
        if date and date.strip() and date.strip() != "Unknown Date":
            source_text += f" | {date}"
            
        if typography_capsule:
            source_spec = typography_capsule.typography_specs.get('source')
            if source_spec:
                create_capsule_based_style(source_para, source_text, source_spec, WD_ALIGN_PARAGRAPH.LEFT)
                logger.info(f"Applied capsule-based source styling: {source_spec.font_family} {source_spec.font_size}pt")
            else:
                create_body_style(source_para, source_text, font_name)
        else:
            create_body_style(source_para, source_text, font_name)
        
        source_path = os.path.join(temp_dir, 'source.docx')
        source_doc.save(source_path)
        components['source'] = source_path
        logger.info(f"Created source document: {source_path}")
    
    # Separate blockquotes and body content
    blockquote_texts = []
    body_texts = []
    
    if structured_content:
        for content_item in structured_content:
            if content_item['type'] == 'blockquote':
                blockquote_texts.append(content_item['text'])
            elif content_item['type'] == 'paragraph':
                body_texts.append({
                    'text': content_item['text'],
                    'indented': content_item.get('indented', False)
                })
    else:
        # Fallback: split regular content by paragraphs
        logger.info(f"INDENT_DEBUG_CONVERTER_B: Using fallback - original text from article_data:\n{article_data.get('text', '')[:500]}...")
        paragraphs = article_data.get('text', '').split('\n\n')
        logger.info(f"INDENT_DEBUG_CONVERTER_B: Split into {len(paragraphs)} paragraphs")
        for i, para in enumerate(paragraphs):
            if para.strip():
                # Check if paragraph starts with spaces (indented)
                is_indented = para.startswith('    ')
                logger.info(f"INDENT_DEBUG_CONVERTER_B: Paragraph {i+1} - indented={is_indented}, text='{para[:60]}...'")
                body_texts.append({'text': para.rstrip(), 'indented': is_indented})
    
    # Create blockquote document
    if blockquote_texts:
        blockquote_doc = Document()
        apply_newspaper_styling(blockquote_doc)
        set_column_layout(blockquote_doc.sections[0], 1)  # Single column
        
        for quote_text in blockquote_texts:
            quote_para = blockquote_doc.add_paragraph()
            
            # Use capsule-based styling if available
            if typography_capsule:
                pullquote_spec = typography_capsule.typography_specs.get('pullquote')
                if pullquote_spec:
                    create_capsule_based_style(quote_para, quote_text, pullquote_spec, WD_ALIGN_PARAGRAPH.JUSTIFY)
                    logger.info(f"Applied capsule-based pullquote styling: {pullquote_spec.font_family} {pullquote_spec.font_size}pt")
                else:
                    create_blockquote_style(quote_para, quote_text, font_name)
            else:
                create_blockquote_style(quote_para, quote_text, font_name)
        
        blockquote_path = os.path.join(temp_dir, 'blockquote.docx')
        blockquote_doc.save(blockquote_path)
        components['blockquote'] = blockquote_path
        logger.info(f"Created blockquote document: {blockquote_path}")
    
    # Create body document
    if body_texts:
        body_doc = Document()
        apply_newspaper_styling(body_doc)
        set_column_layout(body_doc.sections[0], 1)  # Single column
        
        for body_item in body_texts:
            body_para = body_doc.add_paragraph()
            
            # Use capsule-based styling if available
            logger.info(f"INDENT_DEBUG_CONVERTER_C: Creating paragraph - indented={body_item['indented']}, text='{body_item['text'][:60]}...'")
            if typography_capsule:
                body_spec = typography_capsule.typography_specs.get('body')
                if body_spec:
                    create_capsule_based_style(body_para, body_item['text'], body_spec, WD_ALIGN_PARAGRAPH.JUSTIFY, True)
                    logger.info(f"Applied capsule-based body styling: {body_spec.font_family} {body_spec.font_size}pt")
                else:
                    create_body_style(body_para, body_item['text'], font_name, True)
            else:
                create_body_style(body_para, body_item['text'], font_name, True)
        
        body_path = os.path.join(temp_dir, 'body.docx')
        body_doc.save(body_path)
        components['body'] = body_path
        logger.info(f"Created body document: {body_path}")
    
    # Handle images
    downloaded_images = []
    if images:
        logger.info(f"Processing {len(images)} images for article: {article_data.get('headline', 'Unknown')}")
        for i, image in enumerate(images):
            if image.get('is_clipping') and image.get('path'):
                # This is a newspapers.com clipping already saved to disk
                logger.info(f"Using newspapers.com clipping {i+1}: {image['path']}")
                if os.path.exists(image['path']) and os.path.getsize(image['path']) > 0:
                    downloaded_images.append({
                        'path': image['path'],
                        'alt_text': image['alt_text'],
                        'url': image['url']
                    })
                    logger.info(f"Successfully added clipping image {i+1}: {image['path']} ({os.path.getsize(image['path'])} bytes)")
                else:
                    logger.error(f"Clipping image file {image['path']} is missing or empty")
            else:
                # This is a regular image URL that needs to be downloaded
                logger.info(f"Attempting to download image {i+1}: {image['url']}")
                image_path = download_image(image['url'], temp_dir)
                if image_path:
                    # Verify the downloaded file exists and has content
                    if os.path.exists(image_path) and os.path.getsize(image_path) > 0:
                        downloaded_images.append({
                            'path': image_path,
                            'alt_text': image['alt_text'],
                            'url': image['url']
                        })
                        logger.info(f"Successfully downloaded image {i+1}: {image_path} ({os.path.getsize(image_path)} bytes)")
                    else:
                        logger.error(f"Image file {image_path} is missing or empty after download")
                else:
                    logger.error(f"Failed to download image {i+1}: {image['url']}")
    else:
        logger.info(f"No images found for article: {article_data.get('headline', 'Unknown')}")
    
    if downloaded_images:
        components['image'] = downloaded_images
        logger.info(f"Added {len(downloaded_images)} downloaded images to components")
        
        # Create caption document if we have images
        caption_text = f"Image from {article_data.get('source', 'Unknown Source')}"
        if date and date.strip() and date.strip() != "Unknown Date":
            caption_text += f", {date}"
            
        caption_doc = Document()
        apply_newspaper_styling(caption_doc)
        set_column_layout(caption_doc.sections[0], 1)
        
        caption_para = caption_doc.add_paragraph()
        if typography_capsule:
            caption_spec = typography_capsule.typography_specs.get('caption')
            if caption_spec:
                create_capsule_based_style(caption_para, caption_text, caption_spec, WD_ALIGN_PARAGRAPH.CENTER)
                logger.info(f"Applied capsule-based caption styling: {caption_spec.font_family} {caption_spec.font_size}pt")
            else:
                create_body_style(caption_para, caption_text, font_name)
        else:
            create_body_style(caption_para, caption_text, font_name)
        
        caption_path = os.path.join(temp_dir, 'caption.docx')
        caption_doc.save(caption_path)
        components['caption'] = caption_path
        logger.info(f"Created caption document: {caption_path}")
    else:
        logger.warning(f"No images were successfully downloaded for article: {article_data.get('headline', 'Unknown')}")
    
    return components

def create_newspaper_document(md_content, temp_dir):
    """Create a Word document formatted like a newspaper clipping."""
    doc = Document()
    apply_newspaper_styling(doc)
    
    # Extract images from markdown
    images = extract_images_from_markdown(md_content)
    
    # Remove image markdown from content for text processing
    clean_content = md_content
    for image in images:
        clean_content = clean_content.replace(image['markdown'], '')
    
    # Process markdown to get headline and body
    headline, body_text = process_markdown_to_text(clean_content)
    
    # Get font for this content (try to extract URL from markdown if available)
    font_name = get_font_for_site(None)  # Default font since we don't have URL context here
    
    # Force single column layout as specified in requirements
    section = doc.sections[0]
    set_column_layout(section, 1)
    
    # Download images
    downloaded_images = []
    logger.info(f"Starting image download for {len(images)} images")
    
    for i, image in enumerate(images):
        logger.info(f"Processing image {i+1}/{len(images)}: {image['url']}")
        image_path = download_image(image['url'], temp_dir)
        if image_path:
            downloaded_images.append({
                'path': image_path,
                'alt_text': image['alt_text'],
                'url': image['url']  # Keep original URL for zip file processing
            })
            logger.info(f"Successfully processed image {i+1}")
        else:
            logger.warning(f"Failed to download image {i+1}: {image['url']}")
    
    logger.info(f"Downloaded {len(downloaded_images)} out of {len(images)} images")
    
    # Store downloaded images info for zip file processing
    doc._downloaded_images = downloaded_images
    
    # Single column layout with components in order: headline, image, body
    # Add headline
    if headline:
        headline_para = doc.add_paragraph()
        create_headline_style(headline_para, headline, font_name)
    
    # Add image
    if downloaded_images:
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(downloaded_images[0]['path'], width=Inches(3))
        img_para.space_after = Pt(6)
    
    # Add body text with paragraph breaks and indentation preserved
    logger.info(f"INDENT_DEBUG_SIMPLE: Body text for simple conversion:\n{body_text[:500]}...")
    body_paragraphs = body_text.split('\n\n')
    logger.info(f"INDENT_DEBUG_SIMPLE: Split into {len(body_paragraphs)} paragraphs")
    for i, para_text in enumerate(body_paragraphs):
        if para_text.strip():
            body_para = doc.add_paragraph()
            # Check if paragraph starts with spaces (indented)
            is_indented = para_text.startswith('    ')
            logger.info(f"INDENT_DEBUG_SIMPLE: Paragraph {i+1} - indented={is_indented}, text='{para_text[:60]}...'")
            create_body_style(body_para, para_text.rstrip(), font_name, True)
    
    return doc

def convert_markdown_to_newspaper(md_content, output_path=None):
    """Convert markdown content to a newspaper-style Word document."""
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Create the document
        doc = create_newspaper_document(md_content, temp_dir)
        
        # Save the document
        if output_path is None:
            output_path = os.path.join(temp_dir, 'newspaper_clipping.docx')
        
        doc.save(output_path)
        logger.info(f"Newspaper document created: {output_path}")
        
        return output_path
    
    except Exception as e:
        logger.error(f"Error creating newspaper document: {str(e)}")
        raise
    
    finally:
        # Clean up downloaded images
        try:
            import shutil
            shutil.rmtree(temp_dir)
        except Exception as e:
            logger.warning(f"Failed to clean up temp directory: {str(e)}")

def convert_articles_to_component_zip(articles_data, output_zip_path=None):
    """
    Convert articles to separate component Word documents organized by word count, font, and title.
    
    Args:
        articles_data: List of article data dictionaries with structured content
        output_zip_path: Optional path to save zip file
        
    Returns:
        dict: Summary of created files
    """
    temp_dir = tempfile.mkdtemp()
    
    try:
        all_component_files = []
        all_images = []
        article_directories = {}  # Track directories by article
        
        for i, article_data in enumerate(articles_data):
            logger.info(f"Processing article {i+1}/{len(articles_data)}")
            
            # Check if this is a newspapers.com article - treat like normal articles now
            is_newspapers_com = 'newspapers.com' in article_data.get('url', '').lower()
            
            if is_newspapers_com:
                logger.info(f"Processing newspapers.com article normally: {article_data.get('title', 'Unknown')}")
                # Create a simple title component for newspapers.com articles since they have less structured content
                if not article_data.get('structured_content'):
                    # Create minimal structured content for newspapers.com articles
                    article_data['structured_content'] = [
                        {'type': 'paragraph', 'text': f"Directory path where image is stored: article_images/ (exported with Word document)", 'indented': False}
                    ]
                    if article_data.get('headline'):
                        article_data['structured_content'].insert(0, 
                            {'type': 'paragraph', 'text': article_data['headline'], 'indented': False})
                    logger.info("Added structured content for newspapers.com article to show directory path")
            
            # Create component documents for this article
            components = create_component_documents(article_data, temp_dir)
            
            # Get metadata for directory organization
            metadata = components.get('metadata', {})
            directory_name = metadata.get('directory_name', f"article_{i+1}")
            
            # Store directory info
            article_directories[directory_name] = {
                'metadata': metadata,
                'components': [],
                'images': []
            }
            
            # Process each component document
            for component_type, component_path in components.items():
                if component_path and component_type not in ['image', 'metadata']:
                    # Read the document data
                    with open(component_path, 'rb') as f:
                        doc_data = f.read()
                    
                    filename = f"{component_type}.docx"
                    
                    component_info = {
                        'filename': filename,
                        'data': doc_data,
                        'directory': directory_name,
                        'component': component_type,
                        'size': len(doc_data),
                        'metadata': metadata
                    }
                    
                    all_component_files.append(component_info)
                    article_directories[directory_name]['components'].append(component_info)
                    
                    logger.info(f"Created {component_type} document for directory: {directory_name}")
            
            # Handle images
            if components.get('image'):
                logger.info(f"Found {len(components['image'])} images in components for directory: {directory_name}")
                for img_idx, img in enumerate(components['image']):
                    # Make filename more unique by including article index
                    img_filename = f"article_{i+1}_image_{img_idx + 1}_{os.path.basename(img['path'])}"
                    img_info = {
                        'source_path': img['path'],
                        'clip_filename': img_filename,
                        'directory': directory_name,
                        'url': img['url'],
                        'alt_text': img['alt_text']
                    }
                    all_images.append(img_info)
                    article_directories[directory_name]['images'].append(img_info)
                    logger.info(f"Prepared image {img_idx + 1} for zip: {img_filename} from {img['path']}")
            else:
                logger.warning(f"No images found in components for directory: {directory_name}")
                if 'image' in components:
                    logger.warning(f"components['image'] exists but is empty: {components['image']}")
                else:
                    logger.warning(f"components['image'] key not found. Available keys: {list(components.keys())}")
        
        # Create zip file with organized directory structure
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add Word documents organized by article directory
            for doc_file in all_component_files:
                directory_path = doc_file['directory']
                zip_path = f"{directory_path}/{doc_file['filename']}"
                zip_file.writestr(zip_path, doc_file['data'])
                logger.info(f"Added to zip: {zip_path}")
            
            # Add images to their respective article directories
            logger.info(f"Adding {len(all_images)} images to article directories")
            successful_component_images = 0
            for img in all_images:
                try:
                    logger.info(f"Processing component image for zip: {img['clip_filename']} from {img['source_path']}")
                    
                    # Check if source file exists
                    if not os.path.exists(img['source_path']):
                        logger.error(f"Source image file not found: {img['source_path']}")
                        continue
                    
                    # Check file size
                    file_size = os.path.getsize(img['source_path'])
                    if file_size == 0:
                        logger.error(f"Source image file is empty: {img['source_path']}")
                        continue
                    
                    with open(img['source_path'], 'rb') as f:
                        img_data = f.read()
                    
                    directory_path = img['directory']
                    zip_path = f"{directory_path}/images/{img['clip_filename']}"
                    zip_file.writestr(zip_path, img_data)
                    successful_component_images += 1
                    logger.info(f"Successfully added component image to zip: {zip_path} ({len(img_data)} bytes)")
                    
                except Exception as e:
                    logger.error(f"Failed to add image {img['clip_filename']} to zip: {str(e)}")
                    logger.error(f"Image details: source_path={img.get('source_path')}, directory={img.get('directory')}")
            
            logger.info(f"Successfully added {successful_component_images} out of {len(all_images)} component images to zip file")
            
            # Add metadata file for each article directory
            for directory_name, dir_info in article_directories.items():
                metadata = dir_info['metadata']
                metadata_content = f"""Article Information:
Title: {metadata.get('title', 'Unknown')}
Word Count: {metadata.get('word_count', 0)} words
Font Used: {metadata.get('font_name', 'Unknown')}
Components: {len(dir_info['components'])} documents
Images: {len(dir_info['images'])} files
"""
                zip_path = f"{directory_name}/article_info.txt"
                zip_file.writestr(zip_path, metadata_content)
                logger.info(f"Added metadata file: {zip_path}")
        
        zip_data = zip_buffer.getvalue()
        
        # Save zip file if path provided
        if output_zip_path:
            with open(output_zip_path, 'wb') as f:
                f.write(zip_data)
            logger.info(f"Component zip file created: {output_zip_path}")
        
        # Generate summary
        total_size = sum(doc['size'] for doc in all_component_files)
        component_counts = {}
        for doc in all_component_files:
            component_type = doc['component']
            component_counts[component_type] = component_counts.get(component_type, 0) + 1
        
        summary = {
            'zip_data': zip_data,
            'document_count': len(all_component_files),
            'total_size': total_size,
            'image_count': len(all_images),
            'article_count': len(article_directories),
            'component_counts': component_counts,
            'directories': list(article_directories.keys()),
            'articles': [
                {
                    'directory': directory_name,
                    'title': dir_info['metadata'].get('title', 'Unknown'),
                    'word_count': dir_info['metadata'].get('word_count', 0),
                    'font': dir_info['metadata'].get('font_name', 'Unknown'),
                    'components': len(dir_info['components']),
                    'images': len(dir_info['images'])
                } for directory_name, dir_info in article_directories.items()
            ]
        }
        
        logger.info(f"Successfully created organized zip with {len(article_directories)} article directories")
        return summary
    
    except Exception as e:
        logger.error(f"Error creating component zip: {str(e)}")
        raise
    
    finally:
        # Clean up temporary directory
        try:
            import shutil
            shutil.rmtree(temp_dir)
        except Exception as e:
            logger.warning(f"Failed to clean up temp directory: {str(e)}")

def convert_multiple_markdown_to_newspaper_zip(markdown_contents, output_zip_path=None, processed_articles=None):
    """Convert multiple markdown contents to individual Word documents and package in a zip file.
    
    Args:
        markdown_contents: List of markdown strings
        output_zip_path: Optional path to save zip file
        processed_articles: Optional list of article objects with image_data from batch processing
    """
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Create individual Word documents
        doc_files = []
        all_images = []  # Collect all images for the clippings directory
        
        for i, md_content in enumerate(markdown_contents):
            logger.info(f"Processing markdown file {i+1}/{len(markdown_contents)}")
            
            # Extract title for filename
            title = extract_title_from_markdown(md_content)
            safe_filename = sanitize_filename(title)
            
            # Ensure unique filenames
            docx_filename = f"{safe_filename}.docx"
            counter = 1
            original_filename = docx_filename
            while any(doc['filename'] == docx_filename for doc in doc_files):
                docx_filename = f"{safe_filename}-{counter}.docx"
                counter += 1
            
            # Handle processed articles with image_data (from batch processing)
            enhanced_md_content = md_content
            newspapers_image_added = False
            if processed_articles and i < len(processed_articles):
                article = processed_articles[i]
                if article.get('image_data'):
                    logger.info(f"Processing image_data for article: {title}")
                    try:
                        # Save the PIL Image to temp directory with unique identifier
                        image_filename = f"{safe_filename}_article_{i+1}_scraped_image.png"
                        image_path = os.path.join(temp_dir, image_filename)
                        
                        image_saved = False
                        
                        # Convert image_data to PIL Image if it's not already
                        if hasattr(article['image_data'], 'save'):
                            # It's already a PIL Image
                            article['image_data'].save(image_path, 'PNG')
                            image_saved = True
                            logger.info(f"Saved PIL Image for {title}")
                        elif isinstance(article['image_data'], bytes):
                            # It's raw bytes - convert to PIL Image first
                            from PIL import Image
                            img = Image.open(io.BytesIO(article['image_data']))
                            img.save(image_path, 'PNG')
                            image_saved = True
                            logger.info(f"Converted bytes to PIL Image for {title}")
                        elif isinstance(article['image_data'], str):
                            # It might be base64 encoded
                            from PIL import Image
                            import base64
                            img_data = base64.b64decode(article['image_data'])
                            img = Image.open(io.BytesIO(img_data))
                            img.save(image_path, 'PNG')
                            image_saved = True
                            logger.info(f"Converted base64 to PIL Image for {title}")
                        else:
                            logger.warning(f"Unknown image_data format for {title}: {type(article['image_data'])}")
                            logger.warning(f"Image data length: {len(article['image_data']) if hasattr(article['image_data'], '__len__') else 'N/A'}")
                        
                        # Only proceed if image was saved successfully
                        if image_saved:
                            # Add image reference to markdown
                            enhanced_md_content = md_content + f"\n\n![Scraped Article Image]({image_path})\n\n"
                            
                            # Add to clippings collection with unique identifier
                            clip_filename = f"{safe_filename}_article_{i+1}_scraped_image.png"
                            all_images.append({
                                'source_path': image_path,
                                'clip_filename': clip_filename,
                                'url': 'scraped_from_newspapers.com',
                                'alt_text': f'Scraped image for {title}'
                            })
                            newspapers_image_added = True
                            
                            logger.info(f"Successfully processed scraped image for {title} and added to all_images collection")
                            logger.info(f"Total images in collection after newspapers.com addition: {len(all_images)}")
                        
                    except Exception as e:
                        logger.error(f"Failed to process image_data for {title}: {str(e)}")
            
            # Create the document
            doc = create_newspaper_document(enhanced_md_content, temp_dir)
            
            # Collect images from this document (from markdown URLs)
            if hasattr(doc, '_downloaded_images'):
                for img_idx, img in enumerate(doc._downloaded_images):
                    # Create unique filename for the clippings directory
                    img_filename = f"{safe_filename}_article_{i+1}_url_img_{img_idx+1}_{os.path.basename(img['path'])}"
                    all_images.append({
                        'source_path': img['path'],
                        'clip_filename': img_filename,
                        'url': img['url'],
                        'alt_text': img['alt_text']
                    })
                    logger.info(f"Added URL-based image to all_images: {img_filename}")
                    logger.info(f"Total images in collection after URL addition: {len(all_images)}")
            
            # Log final image count for this article
            if newspapers_image_added:
                logger.info(f"Article {i+1} ({title}): Added newspapers.com image to collection")
            elif hasattr(doc, '_downloaded_images') and doc._downloaded_images:
                logger.info(f"Article {i+1} ({title}): Added {len(doc._downloaded_images)} URL-based images to collection")
            else:
                logger.info(f"Article {i+1} ({title}): No images found to add to collection")
            
            # Save to temporary file
            temp_doc_path = os.path.join(temp_dir, docx_filename)
            doc.save(temp_doc_path)
            
            # Read the document data
            with open(temp_doc_path, 'rb') as f:
                doc_data = f.read()
            
            doc_files.append({
                'filename': docx_filename,
                'data': doc_data,
                'title': title,
                'size': len(doc_data)
            })
            
            logger.info(f"Created document: {docx_filename} ({len(doc_data):,} bytes)")
        
        # Create zip file
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add Word documents
            for doc_file in doc_files:
                zip_file.writestr(doc_file['filename'], doc_file['data'])
                logger.info(f"Added to zip: {doc_file['filename']}")
            
            # Add images to clippings directory
            logger.info(f"Adding {len(all_images)} images to clippings directory")
            
            # Debug: Log image sources and check for duplicates
            newspapers_images = len([img for img in all_images if 'scraped_from_newspapers.com' in img['url']])
            url_images = len([img for img in all_images if 'scraped_from_newspapers.com' not in img['url']])
            logger.info(f"Image source breakdown: {newspapers_images} from newspapers.com, {url_images} from URLs")
            
            # Check for duplicate filenames
            clip_filenames = [img['clip_filename'] for img in all_images]
            unique_filenames = set(clip_filenames)
            if len(clip_filenames) != len(unique_filenames):
                logger.warning(f"Found duplicate filenames! {len(clip_filenames)} total vs {len(unique_filenames)} unique")
                # Log duplicate filenames
                from collections import Counter
                filename_counts = Counter(clip_filenames)
                for filename, count in filename_counts.items():
                    if count > 1:
                        logger.warning(f"Duplicate filename: {filename} appears {count} times")
            else:
                logger.info(f"✅ All {len(clip_filenames)} image filenames are unique")
            
            successful_image_count = 0
            for img in all_images:
                try:
                    logger.info(f"Processing image: {img['clip_filename']} from source: {img['source_path']}")
                    
                    # Check if file exists and has content
                    if not os.path.exists(img['source_path']):
                        logger.error(f"Image source file not found: {img['source_path']}")
                        continue
                    
                    file_size = os.path.getsize(img['source_path'])
                    if file_size == 0:
                        logger.error(f"Image source file is empty: {img['source_path']}")
                        continue
                    
                    # Read image data
                    with open(img['source_path'], 'rb') as f:
                        img_data = f.read()
                    
                    # Add to clippings directory in zip
                    zip_path = f"clippings/{img['clip_filename']}"
                    zip_file.writestr(zip_path, img_data)
                    successful_image_count += 1
                    logger.info(f"Successfully added image to zip: {zip_path} ({len(img_data)} bytes)")
                    logger.info(f"Image source: {img.get('url', 'unknown')}, filename: {img['clip_filename']}")
                    
                except Exception as e:
                    logger.error(f"Failed to add image {img['clip_filename']} to zip: {str(e)}")
                    logger.error(f"Image details: source_path={img.get('source_path')}, clip_filename={img.get('clip_filename')}")
            
            logger.info(f"Successfully added {successful_image_count} out of {len(all_images)} images to zip file")
        
        zip_data = zip_buffer.getvalue()
        
        # Save zip file if path provided
        if output_zip_path:
            with open(output_zip_path, 'wb') as f:
                f.write(zip_data)
            logger.info(f"Zip file created: {output_zip_path}")
        
        # Generate summary
        total_size = sum(doc['size'] for doc in doc_files)
        summary = {
            'zip_data': zip_data,
            'document_count': len(doc_files),
            'total_size': total_size,
            'image_count': len(all_images),
            'documents': [{'filename': doc['filename'], 'title': doc['title'], 'size': doc['size']} for doc in doc_files]
        }
        
        logger.info(f"Successfully created zip with {len(doc_files)} newspaper documents and {len(all_images)} images ({total_size:,} bytes total)")
        
        return summary
    
    except Exception as e:
        logger.error(f"Error creating newspaper zip: {str(e)}")
        raise
    
    finally:
        # Clean up temporary directory
        try:
            import shutil
            shutil.rmtree(temp_dir)
        except Exception as e:
            logger.warning(f"Failed to clean up temp directory: {str(e)}")

def create_single_word_document_with_images(articles_data, output_path=None, original_url_order=None, upload_to_drive=True):
    """
    Create a single Word document with all articles in original order and export all images to a separate folder.
    Optionally upload to Google Drive.
    
    Args:
        articles_data: List of article data dictionaries
        output_path: Optional path to save the document
        original_url_order: Optional list of URLs in original document order for preservation
        upload_to_drive: Whether to upload the result to Google Drive (default: True)
        
    Returns:
        dict: Dictionary containing document path, images folder path, summary, and Google Drive info
    """
    logger.info(f"Creating single Word document with {len(articles_data)} articles")
    
    # Preserve original URL order if provided
    if original_url_order:
        logger.info(f"Preserving original URL order from {len(original_url_order)} URLs")
        # Create a URL to article mapping
        url_to_article = {}
        for article in articles_data:
            url = article.get('url', '')
            if url:
                url_to_article[url] = article
        
        # Reorder articles based on original URL order
        ordered_articles = []
        for url in original_url_order:
            if url in url_to_article:
                ordered_articles.append(url_to_article[url])
            else:
                logger.warning(f"URL not found in processed articles: {url}")
        
        # Add any remaining articles that weren't in the original order
        processed_urls = set(original_url_order)
        for article in articles_data:
            url = article.get('url', '')
            if url and url not in processed_urls:
                ordered_articles.append(article)
                logger.info(f"Added article not in original order: {url}")
        
        articles_data = ordered_articles
        logger.info(f"Successfully ordered {len(articles_data)} articles based on original URL order")
    
    # Create main document
    doc = Document()
    apply_newspaper_styling(doc)
    
    # Set single column layout
    section = doc.sections[0]
    set_column_layout(section, 1)
    
    # Create temp directory for processing
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Collections for export
        all_images = []
        processed_articles = []
        
        # Process each article
        for i, article_data in enumerate(articles_data):
            logger.info(f"Processing article {i+1}/{len(articles_data)}: {article_data.get('headline', 'Unknown')}")
            
            # Extract article information
            headline = article_data.get('headline', 'Unknown Article')
            source = article_data.get('source', 'Unknown Source')
            date = article_data.get('date', 'Unknown Date')
            content = article_data.get('full_content') or article_data.get('content', 'No content available')
            url = article_data.get('url', '')
            
            # Get font for this article
            font_name = get_font_for_site(url)
            
            # Add article separator if not the first article
            if i > 0:
                doc.add_page_break()
            
            # Add headline
            headline_para = doc.add_paragraph()
            create_headline_style(headline_para, headline, font_name)
            
            # Add drophead (subheading) if available
            drophead = article_data.get('drophead') or article_data.get('subheading')
            if drophead and drophead.strip():
                drophead_para = doc.add_paragraph()
                drophead_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = drophead_para.add_run(drophead)
                run.font.name = font_name
                run.font.size = Pt(16)
                run.font.italic = True
                drophead_para.space_after = Pt(6)
            
            # Add source and date with enhanced formatting
            if source != 'Unknown Source' or date != 'Unknown Date':
                source_para = doc.add_paragraph()
                source_text = f"Source: {source}"
                if date != 'Unknown Date':
                    source_text += f" | Date: {date}"
                create_body_style(source_para, source_text, font_name, False)
                source_para.space_after = Pt(12)
            
            # Handle article images
            article_images = []
            
            # Check for newspapers.com image_data (PIL Image or bytes)
            if article_data.get('image_data'):
                try:
                    # Enhanced naming: article_{index}_{source}_{headline}.{extension}
                    clean_source = sanitize_filename(source) if source != 'Unknown Source' else 'newspapers'
                    clean_headline = sanitize_filename(headline)
                    image_filename = f"article_{i+1}_{clean_source}_{clean_headline}.png"
                    image_path = os.path.join(temp_dir, image_filename)
                    
                    # Handle different image_data formats
                    if hasattr(article_data['image_data'], 'save'):
                        # It's a PIL Image object
                        article_data['image_data'].save(image_path, 'PNG')
                        logger.info(f"Saved PIL Image from newspapers.com: {image_filename}")
                    elif isinstance(article_data['image_data'], bytes):
                        # It's raw bytes - convert to PIL Image first
                        from PIL import Image
                        import io
                        img = Image.open(io.BytesIO(article_data['image_data']))
                        img.save(image_path, 'PNG')
                        logger.info(f"Converted bytes to PIL Image and saved newspapers.com image: {image_filename}")
                    elif isinstance(article_data['image_data'], str):
                        # It might be base64 encoded
                        from PIL import Image
                        import base64
                        import io
                        img_data = base64.b64decode(article_data['image_data'])
                        img = Image.open(io.BytesIO(img_data))
                        img.save(image_path, 'PNG')
                        logger.info(f"Converted base64 to PIL Image and saved newspapers.com image: {image_filename}")
                    else:
                        logger.error(f"Unknown image_data format: {type(article_data['image_data'])}")
                        continue
                    
                    article_images.append({
                        'path': image_path,
                        'filename': image_filename,
                        'source': 'newspapers.com',
                        'alt_text': f'Newspaper clipping for {headline}'
                    })
                    logger.info(f"Successfully processed newspapers.com image: {image_filename}")
                except Exception as e:
                    logger.error(f"Failed to save newspapers.com image: {str(e)}")
                    logger.error(f"Image data type: {type(article_data.get('image_data'))}")
                    logger.error(f"Image data length: {len(article_data['image_data']) if hasattr(article_data['image_data'], '__len__') else 'N/A'}")
            
            # Check for regular image URL
            elif article_data.get('image_url'):
                try:
                    # Enhanced naming: article_{index}_{source}_{headline}.{extension}
                    clean_source = sanitize_filename(source) if source != 'Unknown Source' else 'web'
                    clean_headline = sanitize_filename(headline)
                    file_ext = os.path.splitext(urlparse(article_data['image_url']).path)[1] or '.jpg'
                    image_filename = f"article_{i+1}_{clean_source}_{clean_headline}{file_ext}"
                    image_path = download_image(article_data['image_url'], temp_dir)
                    
                    if image_path:
                        # Rename to our standardized filename
                        new_image_path = os.path.join(temp_dir, image_filename)
                        os.rename(image_path, new_image_path)
                        
                        article_images.append({
                            'path': new_image_path,
                            'filename': image_filename,
                            'source': 'url',
                            'alt_text': f'Image for {headline}'
                        })
                        logger.info(f"Downloaded image: {image_filename}")
                except Exception as e:
                    logger.error(f"Failed to download image: {str(e)}")
            
            # Add images to document
            if article_images:
                for img in article_images:
                    try:
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        
                        # Add image with appropriate sizing
                        run.add_picture(img['path'], width=Inches(4))
                        img_para.space_after = Pt(6)
                        
                        # Add to collection for export
                        all_images.append(img)
                        logger.info(f"Added {img['source']} image to all_images collection: {img['filename']}")
                        
                    except Exception as e:
                        logger.error(f"Failed to add image to document: {str(e)}")
            
            # Add article content
            if content:
                # Split content into paragraphs
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        content_para = doc.add_paragraph()
                        create_body_style(content_para, para_text.strip(), font_name, True)
            
            # Add directory path information for newspapers.com articles
            is_newspapers_com = 'newspapers.com' in url.lower() if url else False
            if is_newspapers_com and article_images:
                directory_para = doc.add_paragraph()
                directory_text = "Directory path where newspaper image is stored: article_images/ (exported with this document)"
                create_body_style(directory_para, directory_text, font_name, False)
                directory_para.space_after = Pt(12)
            
            # Add URL reference
            if url:
                url_para = doc.add_paragraph()
                create_body_style(url_para, f"URL: {url}", font_name, False)
                url_para.space_after = Pt(18)
            
            processed_articles.append({
                'headline': headline,
                'source': source,
                'date': date,
                'url': url,
                'images': len(article_images)
            })
        
        # Save the document to current directory if no path specified
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"all_articles_{timestamp}.docx"
        
        doc.save(output_path)
        logger.info(f"Single Word document created: {output_path}")
        
        # Create images folder
        images_dir = os.path.join(os.path.dirname(output_path), "article_images")
        os.makedirs(images_dir, exist_ok=True)
        
        # Copy all images to the images folder
        logger.info(f"Processing {len(all_images)} images for export to images folder")
        
        # Debug: Check for duplicate filenames in Word document
        if all_images:
            image_filenames = [img['filename'] for img in all_images]
            unique_filenames = set(image_filenames)
            if len(image_filenames) != len(unique_filenames):
                logger.warning(f"Word doc: Found duplicate image filenames! {len(image_filenames)} total vs {len(unique_filenames)} unique")
                from collections import Counter
                filename_counts = Counter(image_filenames)
                for filename, count in filename_counts.items():
                    if count > 1:
                        logger.warning(f"Word doc: Duplicate filename: {filename} appears {count} times")
            else:
                logger.info(f"Word doc: ✅ All {len(image_filenames)} image filenames are unique")
        
        exported_images = []
        for img in all_images:
            try:
                dest_path = os.path.join(images_dir, img['filename'])
                
                # Copy image to destination
                import shutil
                shutil.copy2(img['path'], dest_path)
                
                exported_images.append({
                    'filename': img['filename'],
                    'source': img['source'],
                    'alt_text': img['alt_text'],
                    'path': dest_path
                })
                
                logger.info(f"Exported {img['source']} image: {img['filename']} to {dest_path}")
            except Exception as e:
                logger.error(f"Failed to export image {img['filename']}: {str(e)}")
        
        # Create base summary
        summary = {
            'document_path': output_path,
            'images_folder': images_dir,
            'articles_count': len(processed_articles),
            'images_count': len(exported_images),
            'articles': processed_articles,
            'images': exported_images,
            'document_size': os.path.getsize(output_path) if os.path.exists(output_path) else 0,
            'google_drive': None
        }
        
        # Upload to Google Drive if requested
        if upload_to_drive:
            logger.info("Attempting to upload to Google Drive...")
            try:
                # Import credential manager to get proper paths
                from utils.credential_manager import CredentialManager
                cred_manager = CredentialManager()
                google_status = cred_manager.get_google_credentials_status()
                
                # Initialize GoogleDriveManager with credential manager paths
                drive_manager = GoogleDriveManager(
                    credentials_path=google_status['credentials_path'],
                    token_path=google_status['token_path'],
                    auto_init=False
                )
                
                # Debug: Log credential paths
                logger.info(f"GoogleDriveManager using credentials: {google_status['credentials_path']}")
                logger.info(f"GoogleDriveManager using token: {google_status['token_path']}")
                logger.info(f"Credentials exist: {google_status['has_credentials']}")
                logger.info(f"Token exists: {google_status['has_token']}")
                
                # Try to initialize from existing credentials first
                init_result = drive_manager.initialize_if_ready()
                
                logger.info(f"GoogleDriveManager init result: {init_result}")
                
                if init_result['success']:
                    # Generate project name based on timestamp and article count
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    project_name = f"Article_Export_{len(processed_articles)}_articles_{timestamp}"
                    
                    # Upload document and images to Google Drive
                    drive_result = drive_manager.upload_document_and_images(
                        document_path=output_path,
                        images_folder_path=images_dir,
                        project_name=project_name
                    )
                    
                    if drive_result['success']:
                        logger.info(f"Successfully uploaded to Google Drive project: {project_name}")
                        summary['google_drive'] = drive_result
                        
                        # Set folder to be publicly accessible
                        try:
                            drive_manager.set_file_permissions(
                                file_id=drive_result['project_folder_id'],
                                role='reader',
                                type='anyone'
                            )
                            logger.info("Set Google Drive folder to be publicly accessible")
                        except Exception as e:
                            logger.warning(f"Failed to set public permissions: {str(e)}")
                    else:
                        logger.error(f"Failed to upload to Google Drive: {drive_result.get('error', 'Unknown error')}")
                        summary['google_drive'] = {'success': False, 'error': drive_result.get('error', 'Upload failed')}
                else:
                    logger.warning(f"Google Drive initialization failed: {init_result.get('error', 'Unknown error')}")
                    
                    # Check if this is a re-authentication requirement (missing refresh token)
                    if init_result.get('requires_reauth'):
                        logger.error("Re-authentication required - missing or invalid refresh token")
                        summary['google_drive'] = {
                            'success': False, 
                            'error': 'Re-authentication required: Please clear Google credentials in sidebar and re-authenticate with offline access',
                            'requires_reauth': True
                        }
                    # If we have credentials but initialization failed, it might be a token issue
                    elif google_status['has_credentials']:
                        logger.info("Attempting full re-authentication since token may be invalid")
                        try:
                            # Try to force a fresh authentication
                            drive_manager._initialize_service()
                            if drive_manager.is_available():
                                logger.info("Successfully authenticated Google Drive with fresh token")
                                
                                # Retry the upload with fresh authentication
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                project_name = f"Article_Export_{len(processed_articles)}_articles_{timestamp}"
                                
                                # Upload document and images to Google Drive
                                drive_result = drive_manager.upload_document_and_images(
                                    document_path=output_path,
                                    images_folder_path=images_dir,
                                    project_name=project_name
                                )
                                
                                if drive_result['success']:
                                    logger.info(f"Successfully uploaded to Google Drive project: {project_name}")
                                    summary['google_drive'] = drive_result
                                    
                                    # Set folder to be publicly accessible
                                    try:
                                        drive_manager.set_file_permissions(
                                            file_id=drive_result['project_folder_id'],
                                            role='reader',
                                            type='anyone'
                                        )
                                        logger.info("Set Google Drive folder to be publicly accessible")
                                    except Exception as e:
                                        logger.warning(f"Failed to set public permissions: {str(e)}")
                                else:
                                    logger.error(f"Upload failed even after re-authentication: {drive_result.get('error', 'Unknown error')}")
                                    summary['google_drive'] = {'success': False, 'error': drive_result.get('error', 'Upload failed after re-auth')}
                            else:
                                logger.error("Re-authentication failed - Google Drive service not available")
                                summary['google_drive'] = {'success': False, 'error': 'Re-authentication failed'}
                        except Exception as reauth_error:
                            logger.error(f"Re-authentication attempt failed: {str(reauth_error)}")
                            summary['google_drive'] = {'success': False, 'error': f'Authentication error: {str(reauth_error)}'}
                    else:
                        logger.warning("Google Drive not configured - no credentials available")
                        summary['google_drive'] = {'success': False, 'error': 'Google Drive not configured'}
                    
            except Exception as e:
                logger.error(f"Google Drive upload failed: {str(e)}")
                summary['google_drive'] = {'success': False, 'error': str(e)}
        
        logger.info(f"Successfully created single Word document with {len(processed_articles)} articles and {len(exported_images)} images")
        image_breakdown = ', '.join([f"{img['source']}({img['filename']})" for img in exported_images])
        logger.info(f"Image breakdown: {image_breakdown}")
        return summary
        
    except Exception as e:
        logger.error(f"Error creating single Word document: {str(e)}")
        raise
    
    finally:
        # Clean up temporary directory
        try:
            import shutil
            shutil.rmtree(temp_dir)
        except Exception as e:
            logger.warning(f"Failed to clean up temp directory: {str(e)}")

def convert_markdown_files_to_newspaper_zip(markdown_file_paths, output_zip_path=None):
    """Convert multiple markdown files to individual Word documents and package in a zip file."""
    markdown_contents = []
    
    for file_path in markdown_file_paths:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                markdown_contents.append(content)
                logger.info(f"Loaded markdown file: {file_path}")
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            raise
    
    if not markdown_contents:
        raise ValueError("No markdown content found in provided files")
    
    return convert_multiple_markdown_to_newspaper_zip(markdown_contents, output_zip_path)

def main():
    """Example usage of the newspaper converter."""
    sample_markdown = """
# Breaking News: AI Revolution Continues

![AI Robot](https://example.com/robot.jpg)

The artificial intelligence industry has seen unprecedented growth this year, with new developments emerging almost daily. Companies across the globe are investing heavily in AI research and development.

This technological revolution is transforming industries from healthcare to finance. Machine learning algorithms are becoming more sophisticated, enabling computers to perform tasks that were once thought to be exclusively human.

Experts predict that the next decade will bring even more dramatic changes as AI systems become more integrated into our daily lives. The implications for society are both exciting and challenging.

As we move forward, it will be crucial to balance innovation with ethical considerations and ensure that the benefits of AI are distributed fairly across all segments of society.
"""
    
    try:
        output_file = convert_markdown_to_newspaper(sample_markdown)
        print(f"Sample newspaper document created: {output_file}")
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    main()