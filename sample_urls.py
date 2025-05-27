#!/usr/bin/env python3
"""
Script to create a sample Word document with URLs for testing the batch processor
"""

from docx import Document
import os

def create_sample_document():
    """Create a sample Word document with various URLs"""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample URLs for Article Extraction', 0)
    
    # Add introduction paragraph
    intro = doc.add_paragraph("""
    This document contains a collection of URLs from various news sources for testing 
    the Best of AI Agent batch processor. The system will extract these URLs and 
    process each one to generate newspaper-style clippings.
    """)
    
    # Add URLs section
    doc.add_heading('News Article URLs', level=1)
    
    # Sample URLs - mix of different sources
    sample_urls = [
        "https://www.espn.com/nfl/story/_/id/38786845/nfl-week-11-takeaways-2023-what-learned-big-questions-every-game-future-team-outlooks",
        "https://www.espn.com/nba/story/_/id/38790123/nba-season-predictions-2024",
        "https://www.espn.com/mlb/story/_/id/38785432/mlb-world-series-recap-2023",
        "https://www.cnn.com/2023/12/01/tech/artificial-intelligence-future/index.html",
        "https://www.bbc.com/news/technology-67890123",
        "https://www.reuters.com/technology/artificial-intelligence-2023-12-01/",
        "https://techcrunch.com/2023/12/01/openai-chatgpt-updates/",
        "https://www.theverge.com/2023/12/1/23987654/tesla-cybertruck-launch",
    ]
    
    # Add URLs as a list
    for i, url in enumerate(sample_urls, 1):
        p = doc.add_paragraph(f"{i}. ")
        p.add_run(url)
    
    # Add section with embedded hyperlinks
    doc.add_heading('Hyperlinked Articles', level=1)
    
    hyperlink_examples = [
        ("ESPN Football Analysis", "https://www.espn.com/nfl/story/_/id/38786845/nfl-week-11-takeaways-2023-what-learned-big-questions-every-game-future-team-outlooks"),
        ("Tech News Update", "https://techcrunch.com/2023/12/01/openai-chatgpt-updates/"),
        ("Sports Analysis", "https://www.espn.com/nba/story/_/id/38790123/nba-season-predictions-2024")
    ]
    
    for title, url in hyperlink_examples:
        p = doc.add_paragraph("• ")
        run = p.add_run(title)
        # Note: Creating actual hyperlinks in python-docx requires more complex code
        # For testing purposes, we'll add the URL in parentheses
        p.add_run(f" ({url})")
    
    # Add table with URLs
    doc.add_heading('URL Table', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Source'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'URL'
    
    # Add data rows
    table_data = [
        ("ESPN", "Sports", "https://www.espn.com/nfl/story/_/id/38786845/nfl-week-11-takeaways-2023-what-learned-big-questions-every-game-future-team-outlooks"),
        ("CNN", "Technology", "https://www.cnn.com/2023/12/01/tech/artificial-intelligence-future/index.html"),
        ("Reuters", "Business", "https://www.reuters.com/technology/artificial-intelligence-2023-12-01/"),
        ("TechCrunch", "Startup", "https://techcrunch.com/2023/12/01/openai-chatgpt-updates/")
    ]
    
    for source, category, url in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = source
        row_cells[1].text = category
        row_cells[2].text = url
    
    # Add footer note
    doc.add_paragraph()
    footer = doc.add_paragraph("""
    Note: This is a sample document for testing purposes. The URLs above are examples 
    and may not all be valid or accessible. The batch processor will attempt to 
    extract content from each URL and generate newspaper-style clippings.
    """)
    footer.italic = True
    
    # Save the document
    output_path = 'sample_urls_document.docx'
    doc.save(output_path)
    
    print(f"Sample document created: {output_path}")
    print(f"Total URLs included: {len(sample_urls) + len(hyperlink_examples) + len(table_data)}")
    
    return output_path

if __name__ == "__main__":
    try:
        output_file = create_sample_document()
        print(f"\nSuccess! Sample document created at: {output_file}")
        print("\nYou can now upload this document to the Streamlit application for testing.")
    except Exception as e:
        print(f"Error creating sample document: {str(e)}")
        print("Make sure you have python-docx installed: pip install python-docx") 